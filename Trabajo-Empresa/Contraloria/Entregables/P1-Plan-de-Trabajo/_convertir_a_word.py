"""
Convierte el Plan de Trabajo (Markdown) a Word, aplicando el formato del Anexo 1 §I del TDR:
- Arial 11
- Espacio simple
- Paginación inferior derecha
- Encabezado del documento

Salida: P1-Plan-de-Trabajo.docx en informe/
"""

import os
import pypandoc
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, "informe", "P1-Plan-de-Trabajo.md")
DOCX_PATH = os.path.join(BASE, "informe", "P1-Plan-de-Trabajo.docx")

# 1) Conversión Markdown → Word con pandoc
pypandoc.convert_file(
    MD_PATH,
    "docx",
    outputfile=DOCX_PATH,
    extra_args=["--standalone"],
)
print(f"[1/3] Pandoc generó: {DOCX_PATH}")

# 2) Post-procesamiento con python-docx
doc = Document(DOCX_PATH)

# Estilo Normal: Arial 11, espacio simple
style_normal = doc.styles["Normal"]
style_normal.font.name = "Arial"
style_normal.font.size = Pt(11)
# rFonts (asegura el East Asian / cs)
rPr = style_normal.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:ascii"), "Arial")
rFonts.set(qn("w:hAnsi"), "Arial")
rFonts.set(qn("w:cs"), "Arial")

# Espaciado simple para todos los párrafos
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rFonts.set(qn("w:cs"), "Arial")

# Aplicar Arial dentro de tablas también
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.append(rFonts)
                    rFonts.set(qn("w:ascii"), "Arial")
                    rFonts.set(qn("w:hAnsi"), "Arial")
                    rFonts.set(qn("w:cs"), "Arial")

print("[2/3] Aplicado Arial 11 + espacio simple")

# 3) Paginación inferior derecha
def add_page_number(paragraph):
    """Inserta el campo PAGE en el párrafo para mostrar el número de página."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Arial"
    run.font.size = Pt(10)


for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # limpiar contenido previo
    for run in list(p.runs):
        run.text = ""
    add_page_number(p)

print("[3/3] Paginación inferior derecha agregada")

doc.save(DOCX_PATH)
print(f"\nWord final: {DOCX_PATH}")
